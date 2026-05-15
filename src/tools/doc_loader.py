# src/tools/doc_loader.py
"""
Universal document loader for user-uploaded files.

Supports:
  .pdf   → LangChain PyPDFLoader (page-by-page, preserves page numbers)
  .docx  → python-docx paragraph extraction
  .txt   → plain text chunked by paragraph
  .csv   → each row becomes a document (useful for data + Q&A)

Returns:
  List of LangChain Document objects with metadata:
    { source: filename, page: N, type: "pdf"|"docx"|"txt"|"csv" }
"""

import os
import io
import csv
import logging
from pathlib import Path

logger = logging.getLogger("atom.doc_loader")


def load_uploaded_file(file_bytes: bytes, filename: str, username: str) -> list:
    """
    Load an uploaded file into LangChain Document objects.

    Args:
        file_bytes : raw bytes of the uploaded file
        filename   : original filename (used to detect type + for citations)
        username   : who uploaded it (stored in metadata)

    Returns:
        List of LangChain Document objects ready for chunking + embedding
    """
    from langchain_core.documents import Document

    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _load_pdf(file_bytes, filename, username)
    elif ext in (".docx", ".doc"):
        return _load_docx(file_bytes, filename, username)
    elif ext == ".txt":
        return _load_txt(file_bytes, filename, username)
    elif ext == ".csv":
        return _load_csv(file_bytes, filename, username)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            "Supported: PDF, DOCX, TXT, CSV"
        )


def _load_pdf(file_bytes: bytes, filename: str, username: str) -> list:
    """Extract text page by page from PDF, with pdfplumber as fallback."""
    from langchain_core.documents import Document
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        # ── Try pdfplumber first (handles more PDF types) ──────────────────
        try:
            import pdfplumber
            docs = []
            with pdfplumber.open(tmp_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    text = text.strip()
                    if text:
                        docs.append(Document(
                            page_content=text,
                            metadata={
                                "source":   filename,
                                "page":     i,
                                "type":     "pdf",
                                "username": username,
                            }
                        ))
            if docs:
                logger.info(f"[DocLoader] PDF '{filename}': {len(docs)} pages via pdfplumber")
                return docs
            logger.warning(f"[DocLoader] pdfplumber found no text in '{filename}', trying PyPDF...")
        except ImportError:
            logger.warning("[DocLoader] pdfplumber not installed, falling back to PyPDFLoader")

        # ── Fallback: PyPDFLoader ──────────────────────────────────────────
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(tmp_path)
        pages  = loader.load()
        docs   = []
        for page in pages:
            content = page.page_content.strip()
            if content:
                docs.append(Document(
                    page_content=content,
                    metadata={
                        "source":   filename,
                        "page":     page.metadata.get("page", 0) + 1,
                        "type":     "pdf",
                        "username": username,
                    }
                ))

        logger.info(f"[DocLoader] PDF '{filename}': {len(docs)} pages via PyPDFLoader")

        if not docs:
            logger.warning(f"[DocLoader] No text extracted from '{filename}' — may be a scanned PDF")

        return docs

    finally:
        os.unlink(tmp_path)


def _load_docx(file_bytes: bytes, filename: str, username: str) -> list:
    """Extract paragraphs from Word document."""
    from langchain_core.documents import Document
    from docx import Document as DocxDocument

    doc   = DocxDocument(io.BytesIO(file_bytes))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Group into ~500-word chunks to simulate pages
    chunks = []
    current_chunk  = []
    current_words  = 0
    chunk_num      = 1

    for para in paras:
        words = len(para.split())
        if current_words + words > 500 and current_chunk:
            chunks.append(Document(
                page_content = "\n\n".join(current_chunk),
                metadata     = {
                    "source":   filename,
                    "page":     chunk_num,
                    "type":     "docx",
                    "username": username,
                }
            ))
            current_chunk = []
            current_words = 0
            chunk_num    += 1
        current_chunk.append(para)
        current_words += words

    if current_chunk:
        chunks.append(Document(
            page_content = "\n\n".join(current_chunk),
            metadata     = {
                "source":   filename,
                "page":     chunk_num,
                "type":     "docx",
                "username": username,
            }
        ))

    logger.info(f"[DocLoader] DOCX '{filename}': {len(chunks)} sections loaded")
    return chunks


def _load_txt(file_bytes: bytes, filename: str, username: str) -> list:
    """Load plain text, split by double newlines (paragraphs)."""
    from langchain_core.documents import Document

    text  = file_bytes.decode("utf-8", errors="replace")
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]

    docs = []
    for i, para in enumerate(paras, 1):
        docs.append(Document(
            page_content = para,
            metadata     = {
                "source":   filename,
                "page":     i,
                "type":     "txt",
                "username": username,
            }
        ))

    logger.info(f"[DocLoader] TXT '{filename}': {len(docs)} paragraphs loaded")
    return docs


def _load_csv(file_bytes: bytes, filename: str, username: str) -> list:
    """Load CSV — each row becomes a document."""
    from langchain_core.documents import Document

    text   = file_bytes.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    docs   = []

    for i, row in enumerate(reader, 1):
        content = "\n".join(f"{k}: {v}" for k, v in row.items() if v and v.strip())
        if content.strip():
            docs.append(Document(
                page_content = content,
                metadata     = {
                    "source":   filename,
                    "page":     i,
                    "type":     "csv",
                    "username": username,
                }
            ))

    logger.info(f"[DocLoader] CSV '{filename}': {len(docs)} rows loaded")
    return docs